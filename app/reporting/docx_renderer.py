from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from app.reporting.models import ScientificReportSnapshot


def _display(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:.6g}"
    return str(value)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)


def _add_table(
    document: Document,
    rows: tuple[dict[str, Any], ...],
    *,
    max_rows: int = 35,
    columns: tuple[str, ...] | None = None,
) -> None:
    if not rows:
        paragraph = document.add_paragraph("Brak danych.")
        paragraph.style = document.styles["Intense Quote"]
        return
    selected_columns = list(columns or tuple(rows[0]))
    table = document.add_table(rows=1, cols=len(selected_columns))
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0].cells
    for index, column in enumerate(selected_columns):
        _set_cell_text(header[index], str(column), bold=True)
        _shade_cell(header[index], "D9E6F2")
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for index, column in enumerate(selected_columns):
            _set_cell_text(cells[index], _display(row.get(column)))
    if len(rows) > max_rows:
        document.add_paragraph(
            f"Tabela skrócona: pokazano {max_rows} z {len(rows)} wierszy. "
            "Pełny zbiór znajduje się w results.xlsx i plikach CSV."
        )


def _add_picture(
    document: Document,
    figures: dict[str, bytes],
    name: str,
    caption: str,
) -> None:
    raw = figures.get(name)
    if raw is None:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(raw), width=Inches(6.4))
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.style = document.styles["Caption"]


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13)):
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Satellite Acquisition Planner — raport automatyczny")


def render_docx(
    snapshot: ScientificReportSnapshot,
    figures: dict[str, bytes],
) -> bytes:
    document = Document()
    _configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(snapshot.title)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(snapshot.project_name).bold = True
    document.add_paragraph(
        f"ID projektu: {snapshot.project_id}\n"
        f"Autor: {snapshot.author or 'nie podano'}\n"
        f"Instytucja: {snapshot.institution or 'nie podano'}\n"
        f"Wygenerowano UTC: {snapshot.generated_at_utc.isoformat()}\n"
        f"Wersja aplikacji: {snapshot.application_version}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if snapshot.description:
        document.add_paragraph(snapshot.description)
    document.add_page_break()

    document.add_heading("1. Podsumowanie", level=1)
    _add_table(document, snapshot.overview_metrics, max_rows=30)
    document.add_paragraph(snapshot.narrative["results"])
    if snapshot.warnings:
        document.add_heading("Uwagi kompletności", level=2)
        for warning in snapshot.warnings:
            document.add_paragraph(warning, style="List Bullet")

    if snapshot.include_methodology:
        document.add_heading("2. Metodyka", level=1)
        document.add_paragraph(snapshot.narrative["methodology"])

    document.add_heading("3. System satelitarny", level=1)
    _add_table(
        document,
        snapshot.satellite_rows,
        max_rows=20,
        columns=(
            "satellite_id",
            "name",
            "sensor_type",
            "altitude_km",
            "inclination_deg",
            "memory_capacity_mb",
            "max_acquisitions_per_day",
        ),
    )

    document.add_heading("4. Zlecenia obserwacyjne", level=1)
    _add_table(
        document,
        snapshot.request_rows,
        max_rows=30,
        columns=(
            "request_id",
            "name",
            "mode",
            "sensor_types",
            "priority",
            "mandatory",
            "earliest_start_utc",
            "latest_end_utc",
        ),
    )

    document.add_heading("5. Okna dostępu i okazje", level=1)
    document.add_heading("5.1. Okna dostępu", level=2)
    _add_table(
        document,
        snapshot.access_rows,
        max_rows=25,
        columns=(
            "window_id",
            "satellite_id",
            "sensor_type",
            "mode_name",
            "start_utc",
            "end_utc",
            "duration_s",
            "coverage_ratio",
        ),
    )
    document.add_heading("5.2. Okazje akwizycyjne", level=2)
    _add_table(
        document,
        snapshot.opportunity_rows,
        max_rows=25,
        columns=(
            "opportunity_id",
            "request_id",
            "satellite_id",
            "sensor_type",
            "start_utc",
            "quality_score",
            "coverage_ratio",
            "is_feasible",
        ),
    )

    document.add_heading("6. Harmonogram", level=1)
    document.add_paragraph(snapshot.narrative["results"])
    _add_picture(
        document,
        figures,
        "schedule_sensor_mix.png",
        "Rysunek 1. Struktura harmonogramu według typu sensora.",
    )
    _add_picture(
        document,
        figures,
        "schedule_satellite_load.png",
        "Rysunek 2. Liczba akwizycji przypisana do satelitów.",
    )
    _add_table(
        document,
        snapshot.schedule_rows,
        max_rows=40,
        columns=(
            "entry_id",
            "request_id",
            "satellite_id",
            "sensor_type",
            "mode_id",
            "start_utc",
            "end_utc",
            "estimated_data_volume_mb",
        ),
    )
    document.add_heading("6.1. Diagnostyka zleceń", level=2)
    _add_picture(
        document,
        figures,
        "unassigned_reasons.png",
        "Rysunek 3. Przyczyny braku pełnej realizacji zleceń.",
    )
    _add_table(
        document,
        snapshot.request_diagnostic_rows,
        max_rows=40,
        columns=(
            "request_id",
            "request_mode",
            "priority",
            "is_mandatory",
            "fulfillment_status",
            "scheduled_entry_count",
            "scheduled_sensor_types",
            "reason_codes",
        ),
    )
    document.add_heading("6.2. Wykorzystanie satelitów", level=2)
    _add_table(
        document,
        snapshot.satellite_kpi_rows,
        max_rows=20,
        columns=(
            "satellite_id",
            "sensor_type",
            "scheduled_acquisitions",
            "imaging_utilization_ratio",
            "memory_utilization_ratio",
            "acquisition_utilization_ratio",
            "generated_data_mb",
        ),
    )

    if snapshot.include_stk_validation:
        document.add_heading("7. Walidacja względem STK", level=1)
        document.add_paragraph(snapshot.narrative["validation"])
        _add_picture(
            document,
            figures,
            "stk_access_errors.png",
            "Rysunek 4. Średnie błędy bezwzględne okien względem STK.",
        )
        _add_table(
            document,
            snapshot.stk_access_rows,
            max_rows=30,
            columns=(
                "model_window_id",
                "stk_interval_id",
                "start_error_s",
                "end_error_s",
                "duration_error_s",
                "overlap_ratio",
            ),
        )
        document.add_heading("7.1. Próbki AER", level=2)
        _add_table(
            document,
            snapshot.stk_aer_rows,
            max_rows=30,
            columns=(
                "timestamp_utc",
                "time_offset_s",
                "azimuth_error_deg",
                "elevation_error_deg",
                "range_error_km",
            ),
        )

    if snapshot.include_benchmarks:
        document.add_heading("8. Benchmark Greedy i CP-SAT", level=1)
        document.add_paragraph(snapshot.narrative["benchmark"])
        _add_picture(
            document,
            figures,
            "benchmark_objective.png",
            "Rysunek 5. Średnia wartość funkcji celu w benchmarku.",
        )
        _add_picture(
            document,
            figures,
            "benchmark_runtime.png",
            "Rysunek 6. Średni czas obliczeń w benchmarku.",
        )
        _add_table(
            document,
            snapshot.benchmark_summary_rows,
            max_rows=40,
            columns=(
                "request_count",
                "algorithm",
                "time_limit_s",
                "run_count",
                "success_count",
                "objective_mean",
                "satisfaction_ratio_mean",
                "runtime_mean_s",
            ),
        )

    document.add_heading("9. Historia harmonogramów", level=1)
    _add_table(
        document,
        snapshot.schedule_history_rows,
        max_rows=40,
        columns=(
            "event_type",
            "algorithm",
            "solver_status",
            "objective_value",
            "fully_satisfied_requests",
            "total_acquisitions",
            "recorded_at_utc",
        ),
    )

    if snapshot.include_limitations:
        document.add_heading("10. Ograniczenia i interpretacja", level=1)
        document.add_paragraph(snapshot.narrative["interpretation"])
        for limitation in snapshot.limitations:
            document.add_paragraph(limitation, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
